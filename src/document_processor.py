"""
document_processor.py
Extracts text from uploaded FIR / Charge Sheet / Witness Statement files
(PDF, Word .docx, or photo/image — including camera captures), classifies
the document type, and pulls out structured facts using an LLM pass.
"""

import re
import os
import json
import pdfplumber
from dataclasses import dataclass, field
from typing import List, Optional
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document as DocxDocument


@dataclass
class CaseDocument:
    doc_type: str = "unknown"          # FIR / Charge Sheet / Witness Statement
    raw_text: str = ""
    fir_number: Optional[str] = None
    police_station: Optional[str] = None
    sections_cited: List[str] = field(default_factory=list)   # e.g. ["IPC 302", "BNS 103"]
    date_of_incident: Optional[str] = None
    date_filed: Optional[str] = None
    complainant: Optional[str] = None
    accused: List[str] = field(default_factory=list)
    witnesses: List[str] = field(default_factory=list)
    allegation_summary: Optional[str] = None
    investigating_officer: Optional[str] = None


class DocumentProcessor:
    """Extracts and structures text from PDF, Word, or image legal documents."""

    SECTION_PATTERN = re.compile(
        r"(?:Section|Sec\.?|U/S|Sections)\s*"
        r"([0-9]{1,4}[A-Za-z]{0,2}(?:\([0-9A-Za-z]{1,4}\))*)"
        r"\s*(?:of\s+the\s+)?"
        r"(IPC|Indian\s+Penal\s+Code|BNS|Bharatiya\s+Nyaya\s+Sanhita|"
        r"CrPC|BNSS|Evidence\s+Act|POCSO|NDPS|IT\s+Act)?",
        re.IGNORECASE,
    )

    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}
    WORD_EXTENSIONS = {".docx"}
    PDF_EXTENSIONS = {".pdf"}

    def extract_text(self, file_path: str) -> str:
        """Dispatches to the right extractor based on file extension.
        Supports: PDF (native + OCR fallback for scans), Word (.docx),
        and photos/images (camera captures or scanned photos) via OCR."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext in self.PDF_EXTENSIONS:
            return self._extract_pdf(file_path)
        elif ext in self.WORD_EXTENSIONS:
            return self._extract_docx(file_path)
        elif ext in self.IMAGE_EXTENSIONS:
            return self._extract_image(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported: PDF, DOCX, and images (PNG/JPG/JPEG)."
            )

    def _extract_pdf(self, pdf_path: str) -> str:
        """Try native text extraction first; fall back to OCR for scanned PDFs."""
        text_chunks = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)

        full_text = "\n".join(text_chunks).strip()

        if len(full_text) < 200:
            full_text = self._ocr_fallback_pdf(pdf_path)

        return full_text

    def _extract_docx(self, docx_path: str) -> str:
        doc = DocxDocument(docx_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()

    def _extract_image(self, image_path: str) -> str:
        """Handles both uploaded photos and camera-captured images via OCR."""
        img = Image.open(image_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        return pytesseract.image_to_string(img, lang="eng").strip()

    def _ocr_fallback_pdf(self, pdf_path: str) -> str:
        images = convert_from_path(pdf_path, dpi=300)
        ocr_text = []
        for img in images:
            ocr_text.append(pytesseract.image_to_string(img, lang="eng"))
        return "\n".join(ocr_text)

    def classify_document(self, text: str) -> str:
        lowered = text.lower()
        if "first information report" in lowered or "fir no" in lowered:
            return "FIR"
        if "charge sheet" in lowered or "final report" in lowered or "chargesheet" in lowered:
            return "Charge Sheet"
        if "witness" in lowered and "statement" in lowered:
            return "Witness Statement"
        if "section 161" in lowered or "161 cr.p.c" in lowered or "161 crpc" in lowered:
            return "Witness Statement"
        return "Unknown"

    def extract_sections(self, text: str) -> List[str]:
        matches = self.SECTION_PATTERN.findall(text)
        sections = set()
        for sec_no, act in matches:
            sec_no_clean = sec_no.strip().strip(",")
            act_clean = (act or "").strip()
            if act_clean:
                sections.add(f"{act_clean} {sec_no_clean}")
            else:
                sections.add(sec_no_clean)
        return sorted(sections)

    def process(self, file_path: str) -> CaseDocument:
        text = self.extract_text(file_path)
        doc_type = self.classify_document(text)
        sections = self.extract_sections(text)

        doc = CaseDocument(
            doc_type=doc_type,
            raw_text=text,
            sections_cited=sections,
        )

        fir_match = re.search(r"FIR\s*No\.?\s*[:\-]?\s*([\w/\-]+)", text, re.IGNORECASE)
        if fir_match:
            doc.fir_number = fir_match.group(1)

        ps_match = re.search(r"Police\s*Station\s*[:\-]?\s*([A-Za-z ]+)", text, re.IGNORECASE)
        if ps_match:
            doc.police_station = ps_match.group(1).strip()

        return doc


def structure_facts_with_llm(llm_client, case_doc: CaseDocument) -> dict:
    """Second pass: LLM extracts facts regex can't reliably catch."""
    system = (
        "You are a legal document parser for Indian criminal case files. "
        "Extract structured facts ONLY from the text given. "
        "Return strict JSON with keys: complainant, accused (list), witnesses (list), "
        "date_of_incident, date_filed, investigating_officer, allegation_summary "
        "(2-4 sentence neutral factual summary, no legal conclusions). "
        "If a field is not present in the text, use null."
    )
    user = f"DOCUMENT TYPE: {case_doc.doc_type}\n\nTEXT:\n{case_doc.raw_text[:12000]}"

    raw_response = llm_client.complete(system, user)
    cleaned = raw_response.strip().strip("```json").strip("```").strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {"error": "Could not parse structured facts", "raw": raw_response}
